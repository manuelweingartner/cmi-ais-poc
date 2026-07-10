# Generates real sample files for the fictional "Staatsarchiv des Kantons
# Dossikon" into public/files/ and writes src/app/data/file-manifest.ts
# with REAL sizes and SHA-512 hashes (fixity, as PREMIS would record it).
#
# Run from repo root:  python scripts/generate_sample_files.py
import hashlib
import io
import json
import os
import zipfile
from datetime import datetime

import fitz  # PyMuPDF
from PIL import Image, ImageDraw
from docx import Document
from openpyxl import Workbook

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "public", "files")
MANIFEST_TS = os.path.join(ROOT, "src", "app", "data", "file-manifest.ts")
os.makedirs(OUT, exist_ok=True)

manifest = []
pid_counter = [0]


def next_pid():
    pid_counter[0] += 1
    return f"STA-AIP-2026-{pid_counter[0]:04d}"


# Honest FIDO-style format identifications (PRONOM naming; PUID only where certain)
FORMATS = {
    "pdf": ("Acrobat PDF 1.7 - Portable Document Format", "1.7", "fmt/276"),
    "docx": ("Microsoft Word for Windows (.docx)", "2007 onwards", "fmt/412"),
    "eml": ("Internet Message Format (E-Mail)", "RFC 5322", None),
    "csv": ("Comma Separated Values", "", "x-fmt/18"),
    "xml": ("Extensible Markup Language", "1.0", "fmt/101"),
    "json": ("JSON Data Interchange Format", "", "fmt/817"),
    "png": ("Portable Network Graphics", "1.2", "fmt/13"),
    "tif": ("Tagged Image File Format", "6.0", "fmt/353"),
    "txt": ("Plain Text File", "", "x-fmt/111"),
    "xlsx": ("Microsoft Excel (.xlsx)", "2007 onwards", "fmt/214"),
    "zip": ("ZIP Format", "", "x-fmt/263"),
}


def add(filename, ve_signatur, ve_titel, titel, events, rendition_of=None, risiko=None):
    path = os.path.join(OUT, filename)
    with open(path, "rb") as f:
        data = f.read()
    ext = filename.rsplit(".", 1)[1].lower()
    fmt_name, fmt_version, puid = FORMATS[ext]
    manifest.append({
        "id": f"f{len(manifest)+1:03d}",
        "fileName": filename,
        "path": f"files/{filename}",
        "extension": ext,
        "sizeBytes": len(data),
        "sha512": hashlib.sha512(data).hexdigest(),
        "formatName": fmt_name,
        "formatVersion": fmt_version,
        "puid": puid,
        "pid": next_pid(),
        "veSignatur": ve_signatur,
        "veTitel": ve_titel,
        "titel": titel,
        "premisEvents": events,
        "renditionOf": rendition_of,
    })


def ev(typ, date, detail):
    return {"typ": typ, "datum": date, "resultat": "erfolgreich", "detail": detail}


EV_STD = lambda d1, d2: [
    ev("Passivierung", d1, "Aus GEVER-Dossier ausgekoppelt, SIP gebildet (eCH-0160)"),
    ev("Ingest", d2, "SIP validiert (Struktur, arelda-Schema, SHA-512), AIP gebildet, PID vergeben"),
]
EV_FIX = ev("Fixity-Pruefung", "2026-06-30", "SHA-512 neu berechnet und mit gespeichertem Digest verglichen")
EV_RET = ev("Retention", "2026-01-15", "Aufbewahrungsregel geprueft, dauernde Archivierung bestaetigt")


# ---------------------------------------------------------------- PDF helpers
def make_pdf(filename, title, lines, header="Staatsarchiv des Kantons Dossikon"):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 60), header, fontsize=9, color=(0.35, 0.39, 0.46))
    page.insert_text((72, 110), title, fontsize=15, color=(0, 0, 0))
    y = 150
    for line in lines:
        page.insert_text((72, y), line, fontsize=10.5, color=(0.1, 0.1, 0.1))
        y += 18
    doc.set_metadata({"title": title, "author": "Staatsarchiv des Kantons Dossikon"})
    doc.save(os.path.join(OUT, filename))
    doc.close()


def make_docx(filename, title, paragraphs):
    d = Document()
    d.add_heading(title, level=1)
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(os.path.join(OUT, filename))


def make_eml(filename, subject, sender, to, date, body):
    content = (
        f"From: {sender}\r\nTo: {to}\r\nSubject: {subject}\r\nDate: {date}\r\n"
        f"MIME-Version: 1.0\r\nContent-Type: text/plain; charset=utf-8\r\n\r\n{body}\r\n"
    )
    with open(os.path.join(OUT, filename), "w", encoding="utf-8", newline="") as f:
        f.write(content)


def make_png(filename, label, size=(480, 320), shade=0):
    im = Image.new("RGB", size, (238 - shade, 234 - shade, 224 - shade))
    dr = ImageDraw.Draw(im)
    for i in range(0, size[0], 24):
        dr.line([(i, 0), (i - 80, size[1])], fill=(220 - shade, 214 - shade, 200 - shade))
    dr.rectangle([12, 12, size[0] - 12, size[1] - 12], outline=(120, 110, 90), width=2)
    dr.text((28, size[1] // 2 - 8), label, fill=(70, 62, 48))
    dr.text((28, size[1] - 34), "Digitalisat - Staatsarchiv Dossikon", fill=(120, 110, 90))
    im.save(os.path.join(OUT, filename))


def make_tif(filename, label, size=(600, 420)):
    im = Image.new("RGB", size, (246, 243, 234))
    dr = ImageDraw.Draw(im)
    for i in range(0, size[0], 40):
        dr.line([(i, 0), (i, size[1])], fill=(228, 222, 205))
    for j in range(0, size[1], 40):
        dr.line([(0, j), (size[0], j)], fill=(228, 222, 205))
    dr.rectangle([10, 10, size[0] - 10, size[1] - 10], outline=(90, 84, 70), width=3)
    dr.text((30, 30), label, fill=(60, 54, 40))
    dr.text((30, size[1] - 40), "Kartensammlung - Digitalisat 600dpi (Ausschnitt)", fill=(110, 100, 80))
    im.save(os.path.join(OUT, filename), format="TIFF")


def make_xlsx(filename, title, headers, rows):
    wb = Workbook()
    ws = wb.active
    ws.title = title[:28]
    ws.append(headers)
    for r in rows:
        ws.append(r)
    wb.save(os.path.join(OUT, filename))


def make_txt(filename, text):
    with open(os.path.join(OUT, filename), "w", encoding="utf-8") as f:
        f.write(text)


def make_csv(filename, headers, rows):
    with open(os.path.join(OUT, filename), "w", encoding="utf-8", newline="") as f:
        f.write(";".join(headers) + "\n")
        for r in rows:
            f.write(";".join(str(x) for x in r) + "\n")


def metadata_xml(ablieferung_id, titel, files):
    items = "\n".join(
        f"      <datei><name>{n}</name><pruefalgorithmus>SHA-512</pruefalgorithmus></datei>" for n in files
    )
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<!-- Vereinfachtes Beispiel eines SIP-Headers in Anlehnung an eCH-0160 (arelda) -->
<paket xmlns="http://bar.admin.ch/arelda/v4" schemaVersion="4.1">
  <paketTyp>SIP</paketTyp>
  <ablieferung>
    <ablieferungsnummer>{ablieferung_id}</ablieferungsnummer>
    <ablieferndeStelle>Staatskanzlei des Kantons Dossikon</ablieferndeStelle>
    <titel>{titel}</titel>
  </ablieferung>
  <inhaltsverzeichnis>
{items}
  </inhaltsverzeichnis>
</paket>
"""


def premis_json(objekt, fmt, events):
    return json.dumps({
        "premis": "3.0",
        "object": objekt,
        "format": fmt,
        "events": events,
    }, ensure_ascii=False, indent=2)


# ================================================================ 1) PDFs (RRB)
RRB = [
    ("A 1.2-004", "Strassenbauprojekt Seetalstrasse", "RRB Nr. 1954/112: Ausbau der Seetalstrasse", "1954-06-22"),
    ("A 1.2-004", "Strassenbauprojekt Seetalstrasse", "RRB Nr. 1954/171: Landerwerb Seetalstrasse", "1954-09-14"),
    ("A 1.2-007", "Kantonsspital Erweiterungsbau", "RRB Nr. 1961/043: Projektierungskredit Kantonsspital", "1961-03-07"),
    ("A 1.2-007", "Kantonsspital Erweiterungsbau", "RRB Nr. 1962/205: Baukredit Erweiterungsbau", "1962-11-20"),
    ("A 1.3-001", "Gesamterneuerungswahlen 1978", "RRB Nr. 1978/002: Wahlanordnung Gesamterneuerungswahlen", "1978-01-10"),
    ("A 1.3-012", "Kantonale Abstimmungen 1985", "RRB Nr. 1985/077: Festsetzung Abstimmungstermine", "1985-04-16"),
    ("A 1.4-003", "Schulhausbauten Gemeinde Lindberg", "RRB Nr. 1990/154: Subvention Schulhaus Lindberg", "1990-08-28"),
    ("A 1.4-009", "Hochwasserschutz Dossibach", "RRB Nr. 1999/231: Sofortmassnahmen Hochwasserschutz", "1999-05-25"),
    ("A 1.5-002", "Verwaltungsreorganisation 2004", "RRB Nr. 2004/018: Zusammenlegung der Aemter", "2004-02-03"),
    ("A 1.5-011", "E-Government-Strategie", "RRB Nr. 2008/122: E-Government-Strategie des Kantons", "2008-06-17"),
]
for i, (sig, dossier, titel, datum) in enumerate(RRB, 1):
    fn = f"rrb_{datum[:4]}_{i:03d}.pdf"
    make_pdf(fn, titel, [
        f"Beschluss des Regierungsrates vom {datum}",
        "",
        "Der Regierungsrat des Kantons Dossikon,",
        "gestuetzt auf die Verfassung des Kantons Dossikon,",
        "auf Antrag der zustaendigen Direktion,",
        "",
        "beschliesst:",
        "",
        "1. Dem Geschaeft wird zugestimmt.",
        "2. Die Direktion wird mit dem Vollzug beauftragt.",
        "3. Mitteilung an die Beteiligten.",
        "",
        "Im Namen des Regierungsrates",
        "Der Staatsschreiber",
    ])
    add(fn, sig, dossier, titel, EV_STD(datum[:4] + "-12-31", "2025-11-12") + [EV_FIX])

# Urteile Obergericht (PDF)
URTEILE = [
    ("B 1.1-003", "Zivilrechtliche Berufungen 1995", "Urteil OG 1995/44 (anonymisiert)", "1995-10-04"),
    ("B 1.1-003", "Zivilrechtliche Berufungen 1995", "Urteil OG 1995/61 (anonymisiert)", "1995-12-12"),
    ("B 1.2-001", "Verwaltungsgerichtsbeschwerden 2001", "Urteil VG 2001/17 (anonymisiert)", "2001-06-08"),
]
for i, (sig, dossier, titel, datum) in enumerate(URTEILE, 1):
    fn = f"urteil_{datum[:4]}_{i:02d}.pdf"
    make_pdf(fn, titel, [
        f"Urteil vom {datum}",
        "",
        "In Sachen A. gegen B. betreffend Forderung.",
        "Die Parteien sind zur Wahrung des Persoenlichkeitsschutzes anonymisiert.",
        "",
        "Das Gericht erkennt:",
        "1. Die Berufung wird abgewiesen.",
        "2. Die Kosten werden der berufungsfuehrenden Partei auferlegt.",
    ], header="Obergericht des Kantons Dossikon")
    add(fn, sig, dossier, titel, EV_STD(datum[:4] + "-12-31", "2025-11-18"))

# Ablieferungsquittungen (PDF) - used by Datenuebernahme app
for i, (ablnr, titel) in enumerate([
    ("ABL-2026-003", "Ablieferung Staatskanzlei: RRB-Serie 2000-2010"),
    ("ABL-2026-005", "Ablieferung Baudirektion: Projektdossiers Umfahrung Lindberg"),
], 1):
    fn = f"ablieferungsquittung_{ablnr.lower().replace('-', '_')}.pdf"
    make_pdf(fn, f"Ablieferungsquittung {ablnr}", [
        "Uebernahmenachweis gemaess Ingest-Prozess",
        "",
        f"Ablieferung: {titel}",
        "Abliefernde Stelle: siehe Ablieferungsdossier",
        "",
        "Die Ablieferung wurde entgegengenommen und validiert:",
        "- Ordnerstruktur geprueft (header/metadata.xml vorhanden)",
        "- metadata.xml gegen arelda-Schema (eCH-0160) validiert",
        "- SHA-512-Pruefsummen aller Dateien verifiziert",
        "",
        "Die Uebernahme ist damit quittiert.",
        "",
        "Staatsarchiv des Kantons Dossikon, Bereich Ingest",
    ])
    add(fn, "ABL", "Ablieferungen", f"Ablieferungsquittung {ablnr}", [ev("Ingest", "2026-05-0%d" % i, "Quittung nach erfolgreicher Uebernahme erzeugt")])

# Bauplaene als PDF
for i, (sig, dossier, titel) in enumerate([
    ("A 2.1-005", "Umfahrung Lindberg, Projektierung", "Situationsplan Umfahrung Lindberg 1:5000"),
    ("A 2.1-005", "Umfahrung Lindberg, Projektierung", "Laengsprofil Umfahrung Lindberg"),
    ("A 2.2-002", "Kantonsschule Neubau 1972", "Grundrissplan Erdgeschoss Kantonsschule"),
], 1):
    fn = f"plan_{i:02d}.pdf"
    make_pdf(fn, titel, [
        "Planbeilage zum Projektdossier (Reproduktion)",
        "",
        "Massstab und Legende gemaess Originalplan.",
        "Original im Planarchiv, Digitalisat als Arbeitskopie.",
    ], header="Baudirektion des Kantons Dossikon")
    add(fn, sig, dossier, titel, EV_STD("2019-03-30", "2025-12-02"))

# ================================================================ 2) DOCX (+ PDF-Renditionen)
DOCX_ITEMS = [
    ("A 1.5-011", "E-Government-Strategie", "Projektbericht E-Government Phase 1",
     ["Der vorliegende Bericht fasst die Ergebnisse der Phase 1 zusammen.",
      "Die Pilotdienste wurden erfolgreich eingefuehrt.",
      "Empfehlung: Ausbau der elektronischen Schalterdienste."]),
    ("A 2.1-005", "Umfahrung Lindberg, Projektierung", "Technischer Bericht Umfahrung Lindberg",
     ["Der technische Bericht beschreibt Linienfuehrung und Kunstbauten.",
      "Die Bruecke ueber den Dossibach ist als Bogenbruecke konzipiert."]),
    ("A 3.1-002", "Kantonsspital Jahresberichte", "Jahresbericht Kantonsspital 1998",
     ["Das Berichtsjahr war gepraegt von der Eroeffnung des Neubaus.",
      "Die Patientenzahlen stiegen um 4.2 Prozent."]),
    ("C 1.1-004", "Korrespondenz Familie Steinmann", "Brief an Regierungsrat Steinmann (Transkript)",
     ["Sehr geehrter Herr Regierungsrat",
      "Mit diesem Schreiben erlauben wir uns, Ihnen fuer Ihr Wirken zu danken.",
      "Hochachtungsvoll, die Buergergemeinde Lindberg"]),
    ("A 1.5-002", "Verwaltungsreorganisation 2004", "Schlussbericht Verwaltungsreorganisation",
     ["Die Reorganisation wurde per 1. Januar 2005 abgeschlossen.",
      "Die Zusammenlegung der Aemter fuehrte zu Effizienzgewinnen."]),
    ("B 2.1-001", "Notariatskreis Dossikon-Stadt", "Weisung an die Notariate 1988 (Abschrift)",
     ["Die Notariate werden angewiesen, die Register nach neuem Muster zu fuehren.",
      "Die Umstellung hat bis Ende 1988 zu erfolgen."]),
    ("D 2.1-003", "Karten und Plaene, Erschliessungsnotizen", "Erschliessungsnotiz Kartensammlung",
     ["Die Kartensammlung umfasst rund 4500 Blaetter.",
      "Die Erschliessung erfolgt nach Regionen und Epochen."]),
]
for i, (sig, dossier, titel, paras) in enumerate(DOCX_ITEMS, 1):
    fn = f"dok_{i:02d}.docx"
    make_docx(fn, titel, paras)
    add(fn, sig, dossier, titel, EV_STD("2020-06-30", "2025-11-25"))

# PDF/A-artige Renditionen fuer 3 DOCX (ehrlich als PDF 1.7 deklariert, Rendition-Beziehung)
for i in (1, 2, 5):
    src = manifest[[m["fileName"] for m in manifest].index(f"dok_{i:02d}.docx")]
    fn = f"dok_{i:02d}_rendition.pdf"
    make_pdf(fn, src["titel"], ["Archivische Nutzungsform (Rendition) des Originals " + src["fileName"], "",
                                "Inhalt siehe Original; erzeugt durch den Document Service."])
    add(fn, src["veSignatur"], src["veTitel"], src["titel"] + " (Rendition)",
        [ev("Formatmigration", "2025-11-26", f"Rendition aus {src['fileName']} erzeugt (Document Service)")],
        rendition_of=src["id"])

# ================================================================ 3) EML
EMLS = [
    ("A 2.1-005", "Umfahrung Lindberg, Projektierung", "Stellungnahme Gemeinde Lindberg zur Linienfuehrung",
     "gemeindekanzlei@lindberg.dk.ch", "baudirektion@dk.ch", "Mon, 12 Mar 2018 09:14:00 +0100",
     "Sehr geehrte Damen und Herren\n\nDie Gemeinde Lindberg unterstuetzt die Variante Ost.\n\nFreundliche Gruesse\nGemeindekanzlei Lindberg"),
    ("A 2.1-005", "Umfahrung Lindberg, Projektierung", "RE: Stellungnahme Gemeinde Lindberg",
     "baudirektion@dk.ch", "gemeindekanzlei@lindberg.dk.ch", "Tue, 13 Mar 2018 14:02:00 +0100",
     "Guten Tag\n\nBesten Dank fuer die Stellungnahme. Sie wird ins Mitwirkungsverfahren aufgenommen.\n\nBaudirektion Kanton Dossikon"),
    ("A 1.5-011", "E-Government-Strategie", "Projektfreigabe Phase 2",
     "staatskanzlei@dk.ch", "projektleitung.egov@dk.ch", "Thu, 05 Feb 2009 11:30:00 +0100",
     "Der Regierungsrat hat die Phase 2 freigegeben. Details im RRB 2009/031."),
    ("A 3.1-002", "Kantonsspital Jahresberichte", "Zustellung Jahresbericht 1998",
     "direktion@ksdk.ch", "gesundheitsdirektion@dk.ch", "Fri, 26 Mar 1999 08:45:00 +0100",
     "In der Beilage der Jahresbericht 1998 zuhanden der Direktion."),
    ("B 2.1-001", "Notariatskreis Dossikon-Stadt", "Rueckfrage Registerfuehrung",
     "notariat.stadt@dk.ch", "obergericht@dk.ch", "Wed, 17 Aug 1988 10:12:00 +0100",
     "Wir bitten um Praezisierung der Weisung betreffend Registerfuehrung, Ziffer 3."),
]
for i, (sig, dossier, subj, frm, to, date, body) in enumerate(EMLS, 1):
    fn = f"mail_{i:02d}.eml"
    make_eml(fn, subj, frm, to, date, body)
    add(fn, sig, dossier, subj, EV_STD("2021-01-31", "2025-12-05"), risiko=None)

# ================================================================ 4) CSV / XLSX
make_csv("statistik_einwohner.csv", ["Jahr", "Einwohner", "Gemeinden"],
         [(1950, 98230, 74), (1970, 121450, 74), (1990, 138920, 71), (2010, 152300, 68), (2020, 159840, 65)])
add("statistik_einwohner.csv", "A 4.1-001", "Kantonale Statistik, Einwohnerzahlen", "Einwohnerstatistik 1950-2020",
    EV_STD("2022-02-28", "2025-12-05"))

make_csv("pegel_falkenbach.csv", ["Datum", "PegelstandCm"],
         [("1999-05-2%d" % d, 180 + d * 22) for d in range(1, 6)])
add("pegel_falkenbach.csv", "A 1.4-009", "Hochwasserschutz Dossibach", "Pegelmessungen Dossibach Mai 1999",
    EV_STD("2005-12-31", "2025-12-05"))

make_csv("bestandesliste_fotos.csv", ["Signatur", "Titel", "Jahr"],
         [("D 1.1-%03d" % i, "Fotografie Ortsansicht Nr. %d" % i, 1920 + i * 7) for i in range(1, 7)])
add("bestandesliste_fotos.csv", "D 1.1", "Fotosammlung", "Bestandesliste Fotosammlung (Auszug)",
    [ev("Ingest", "2025-12-05", "Erschliessungshilfe bei Uebernahme mitgeliefert")])

make_csv("kassationsliste_2026.csv", ["Signatur", "Titel", "Entscheid"],
         [("A 9.1-001", "Doppel Kreditorenbelege 1998", "kassiert"), ("A 9.1-002", "Sitzungseinladungen 2001", "kassiert")])
add("kassationsliste_2026.csv", "A 9.1", "Bewertung und Kassation", "Kassationsliste 2026 (Auszug)",
    [ev("Ingest", "2026-02-10", "Bewertungsentscheid dokumentiert")])

make_xlsx("steuerstatistik_2005.xlsx", "Steuerstatistik", ["Gemeinde", "Steuerertrag CHF"],
          [["Dossikon-Stadt", 48200000], ["Lindberg", 9300000], ["Seetal", 6100000]])
add("steuerstatistik_2005.xlsx", "A 4.1-004", "Kantonale Statistik, Steuern", "Steuerstatistik 2005",
    EV_STD("2015-12-31", "2025-12-06"))

make_xlsx("magazinbelegung_2026.xlsx", "Magazinbelegung", ["Standort", "Laufmeter belegt", "Laufmeter total"],
          [["A", 1240, 1600], ["A.01", 420, 500], ["B", 890, 1200]])
add("magazinbelegung_2026.xlsx", "MAG", "Magazinverwaltung", "Magazinbelegung 2026 (Arbeitsliste)",
    [ev("Ingest", "2026-03-01", "Interne Arbeitsliste uebernommen")])

make_xlsx("akzessionsjournal_2025.xlsx", "Akzessionsjournal", ["Akzessionsnummer", "Titel", "Status"],
          [["AKZ 2025/01", "Uebernahme Staatskanzlei RRB", "Abgeschlossen"],
           ["AKZ 2025/07", "Nachlass Prof. Huber", "InBearbeitung"]])
add("akzessionsjournal_2025.xlsx", "AKZ", "Akzessionen", "Akzessionsjournal 2025 (Auszug)",
    [ev("Ingest", "2026-01-08", "Journal als Arbeitsunterlage uebernommen")])

# ================================================================ 5) XML / JSON
xml1 = metadata_xml("ABL-2026-003", "RRB-Serie 2000-2010", ["rrb_2004_009.pdf", "rrb_2008_010.pdf"])
make_txt("metadata_abl_2026_003.xml", xml1)
add("metadata_abl_2026_003.xml", "ABL-2026-003", "Ablieferung Staatskanzlei", "SIP-Header metadata.xml (ABL-2026-003)",
    [ev("Ingest", "2026-05-01", "Gegen arelda-Schema validiert")])

xml2 = metadata_xml("ABL-2026-005", "Projektdossiers Umfahrung Lindberg", ["plan_01.pdf", "plan_02.pdf", "mail_01.eml"])
make_txt("metadata_abl_2026_005.xml", xml2)
add("metadata_abl_2026_005.xml", "ABL-2026-005", "Ablieferung Baudirektion", "SIP-Header metadata.xml (ABL-2026-005)",
    [ev("Ingest", "2026-05-02", "Gegen arelda-Schema validiert")])

make_txt("ead_export_bestand_a1.xml", """<?xml version="1.0" encoding="UTF-8"?>
<!-- Vereinfachter EAD-Export (Auszug) des Bestandes A 1 -->
<ead xmlns="urn:isbn:1-931666-22-9">
  <archdesc level="fonds">
    <did>
      <unitid>A 1</unitid>
      <unittitle>Regierungsratsbeschluesse (RRB)</unittitle>
      <unitdate>1850-2010</unitdate>
    </did>
  </archdesc>
</ead>
""")
add("ead_export_bestand_a1.xml", "A 1", "Regierungsratsbeschluesse (RRB)", "EAD-Export Bestand A 1 (Auszug)",
    [ev("Ingest", "2026-04-14", "EAD-Export erzeugt (Dissemination)")])

# PREMIS JSON files (per slide: PREMIS als JSON je Ereignis/Objekt)
for i, src_idx in enumerate([0, 10, 15], 1):
    src = manifest[src_idx]
    fn = f"premis_{src['id']}.json"
    make_txt(fn, premis_json(
        {"pid": src["pid"], "originalName": src["fileName"], "sizeBytes": src["sizeBytes"],
         "fixity": {"algorithm": "SHA-512", "digest": src["sha512"][:32] + "..."}},
        {"name": src["formatName"], "version": src["formatVersion"], "registry": "PRONOM", "puid": src["puid"]},
        src["premisEvents"]))
    add(fn, src["veSignatur"], src["veTitel"], f"PREMIS-Metadaten zu {src['fileName']}",
        [ev("Ingest", "2025-12-06", "PREMIS bei Uebernahme erzeugt (FIDO-Formaterkennung)")])

# ================================================================ 6) PNG / TIF
FOTOS = [
    ("D 1.1-001", "Fotosammlung", "Ortsansicht Dossikon um 1920"),
    ("D 1.1-002", "Fotosammlung", "Marktplatz Dossikon 1935"),
    ("D 1.1-003", "Fotosammlung", "Bahnhofquartier 1948"),
    ("D 1.1-004", "Fotosammlung", "Hochwasser Dossibach 1999"),
    ("D 1.1-005", "Fotosammlung", "Luftaufnahme Kantonsschule 1975"),
    ("D 1.1-006", "Fotosammlung", "Eroeffnung Umfahrung Lindberg 2023"),
]
for i, (sig, dossier, titel) in enumerate(FOTOS, 1):
    fn = f"foto_{i:02d}.png"
    make_png(fn, titel, shade=i * 6)
    add(fn, sig, dossier, titel, EV_STD("2024-08-31", "2025-12-08") + [EV_RET])

KARTEN = [
    ("D 2.1-001", "Karten und Plaene", "Kantonskarte Dossikon 1878 (Ausschnitt)"),
    ("D 2.1-002", "Karten und Plaene", "Uebersichtsplan Dossikon-Stadt 1902"),
    ("D 2.1-004", "Karten und Plaene", "Gewaesserkarte Dossibachtal 1955"),
]
for i, (sig, dossier, titel) in enumerate(KARTEN, 1):
    fn = f"karte_{i:02d}.tif"
    make_tif(fn, titel)
    add(fn, sig, dossier, titel, EV_STD("2023-05-31", "2025-12-08"))

# ================================================================ 7) TXT
make_txt("transkript_ratsprotokoll_1852.txt",
         "Transkription Ratsprotokoll vom 14. Herbstmonat 1852 (Auszug)\n\n"
         "Der Rath beschliesst, die Strasse gegen Lindberg auszubessern\n"
         "und hierfuer die Summe von 1200 Franken zu bewilligen.\n")
add("transkript_ratsprotokoll_1852.txt", "A 1.1-001", "Ratsprotokolle 1850-1899",
    "Transkription Ratsprotokoll 1852 (Auszug)", EV_STD("2016-01-31", "2025-12-08"))

make_txt("findmittel_hinweis_c1.txt",
         "Findmittel-Hinweis Familienarchiv Steinmann\n\n"
         "Das Familienarchiv wurde 2019 als Depositum uebernommen.\n"
         "Benutzung nur mit schriftlicher Bewilligung der Familie.\n")
add("findmittel_hinweis_c1.txt", "C 1.1", "Familienarchiv Steinmann", "Findmittel-Hinweis Familienarchiv",
    [ev("Ingest", "2025-12-08", "Begleitdokument der Uebernahme")])

make_txt("lesehilfe_notariatsregister.txt",
         "Lesehilfe Notariatsregister 1860-1920\n\n"
         "Die Registerbaende sind chronologisch gefuehrt.\n"
         "Abkuerzungen: V. = Verkauf, T. = Tausch, E. = Erbteilung.\n")
add("lesehilfe_notariatsregister.txt", "B 2.1", "Notariatsregister", "Lesehilfe Notariatsregister",
    [ev("Ingest", "2025-12-08", "Begleitdokument der Uebernahme")])

# ================================================================ 8) SIP ZIPs
def make_sip_zip(filename, ablnr, titel, inner):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("header/metadata.xml", metadata_xml(ablnr, titel, [n for n, _ in inner]))
        for name, content in inner:
            z.writestr(f"content/{name}", content)
    with open(os.path.join(OUT, filename), "wb") as f:
        f.write(buf.getvalue())

make_sip_zip("SIP_20260430_STA_1.zip", "ABL-2026-003", "RRB-Serie 2000-2010",
             [("beschlussliste.txt", "RRB 2000-2010, Liste der Beschluesse (Beilage)")])
add("SIP_20260430_STA_1.zip", "ABL-2026-003", "Ablieferung Staatskanzlei", "SIP-Paket ABL-2026-003",
    [ev("Ingest", "2026-05-01", "SIP entgegengenommen, Struktur und Fixity geprueft")])

make_sip_zip("SIP_20260502_BAU_1.zip", "ABL-2026-005", "Projektdossiers Umfahrung Lindberg",
             [("dossierliste.txt", "Projektdossiers Umfahrung Lindberg, Uebersicht")])
add("SIP_20260502_BAU_1.zip", "ABL-2026-005", "Ablieferung Baudirektion", "SIP-Paket ABL-2026-005",
    [ev("Ingest", "2026-05-02", "SIP entgegengenommen, Struktur und Fixity geprueft")])

# Aeltere Office-Formate als Migrationskandidaten gibt es hier bewusst NICHT als
# echte Dateien (kein ehrliches .doc/.msg erzeugbar) - Formatrisiken werden in der
# App auf Basis der vorhandenen Formate dargestellt.

# ================================================================ Manifest schreiben
manifest.sort(key=lambda m: (m["extension"], m["fileName"]))
# re-id after sort? keep original ids stable instead: do NOT re-id (renditionOf references)

lines = []
lines.append("// GENERATED by scripts/generate_sample_files.py - do not edit manually.")
lines.append("// Real files live in public/files/; sizes and SHA-512 digests are computed")
lines.append("// from the actual file contents (honest fixity values).")
lines.append("import { ArchivDatei } from '../models/ais.model';")
lines.append("")
lines.append("export const FILE_MANIFEST: ArchivDatei[] = [")
for m in manifest:
    events = ", ".join(
        "{ typ: '%s', datum: '%s', resultat: '%s', detail: '%s' }" % (e["typ"], e["datum"], e["resultat"], e["detail"].replace("'", "\\'"))
        for e in m["premisEvents"])
    puid = f"'{m['puid']}'" if m["puid"] else "undefined"
    rend = f"'{m['renditionOf']}'" if m.get("renditionOf") else "undefined"
    lines.append(
        "  { id: '%s', fileName: '%s', path: '%s', extension: '%s', sizeBytes: %d, sha512: '%s', "
        "formatName: '%s', formatVersion: '%s', puid: %s, pid: '%s', veSignatur: '%s', veTitel: '%s', "
        "titel: '%s', premisEvents: [%s], renditionOf: %s }," % (
            m["id"], m["fileName"], m["path"], m["extension"], m["sizeBytes"], m["sha512"],
            m["formatName"].replace("'", "\\'"), m["formatVersion"], puid, m["pid"], m["veSignatur"],
            m["veTitel"].replace("'", "\\'"), m["titel"].replace("'", "\\'"), events, rend))
lines.append("];")
lines.append("")

os.makedirs(os.path.dirname(MANIFEST_TS), exist_ok=True)
with open(MANIFEST_TS, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

from collections import Counter
c = Counter(m["extension"] for m in manifest)
total = sum(m["sizeBytes"] for m in manifest)
print(f"{len(manifest)} files, {total/1024:.0f} KB total")
print(dict(sorted(c.items())))
print("manifest written:", MANIFEST_TS)
